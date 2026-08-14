from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "Инструкция_пользователя_Система_командировок.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = RGBColor(31, 37, 48)
MUTED = RGBColor(92, 99, 112)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_number(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.bold = True
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Система планирования командировок"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Страница ").font.size = Pt(9)
    add_page_field(footer)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Инструкция пользователя")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    r = subtitle.add_run("Система планирования командировок")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED

    add_note(doc, "Назначение", "Приложение используется для планирования командировок, их согласования по подчинённости, просмотра календаря и рабочих коммуникаций.")
    add_heading(doc, "Краткий старт", 1)
    for item in [
        "Откройте адрес приложения, выданный компанией, в современном браузере.",
        "Введите рабочий email и пароль, полученные от администратора.",
        "Проверьте в разделе «Мой профиль» свои ФИО, отдел и руководителя. От этих сведений зависят видимость коллег, маршрутизация согласования и список собеседников в чате.",
        "Для создания новой заявки перейдите в «Мои командировки» и нажмите «Создать командировку».",
    ]:
        add_number(doc, item)
    add_note(doc, "Важно", "Не передавайте пароль другим лицам. Если вход не выполняется или указаны неверные ФИО, отдел либо руководитель, используйте «Мой профиль» -> «Связь с администратором».")

    add_heading(doc, "1. Роли и доступы", 1)
    doc.add_paragraph("Доступ к данным определяется вашей учётной записью. Сотрудник видит свои рабочие данные; руководители дополнительно получают заявки подчинённых для согласования; администратор управляет справочниками и настройками.")
    add_table(doc, ["Роль", "Основные действия"], [
        ("Сотрудник", "Создаёт и отслеживает свои командировки, смотрит календарь, использует чат, профиль и обращение к администратору."),
        ("Руководитель", "Имеет возможности сотрудника и раздел «Согласование» для заявок в своей зоне ответственности."),
        ("Администратор", "Видит расширенные данные, завершает согласование при необходимости, ведёт пользователей, города, маршруты, праздники, суточные и служебные настройки."),
    ], [2700, 6660])
    add_note(doc, "Зависимости", "Руководитель и отдел назначаются администратором. Если эти данные не заполнены или неверны, заявка может попасть не тому согласующему, а список контактов в чате будет неполным.")

    add_heading(doc, "2. Создание командировки", 1)
    doc.add_paragraph("В разделе «Мои командировки» отображается история ваших заявок. Для новой заявки откройте форму создания и последовательно заполните обязательные поля.")
    add_table(doc, ["Поле", "Что указать"], [
        ("Маршрут", "Начните вводить название и выберите маршрут из выпадающего списка. Маршруты ведутся в общем справочнике."),
        ("Вид транспорта", "Выберите автомобиль, железнодорожный или авиационный транспорт."),
        ("Дата начала и окончания", "Выберите даты в календаре. Дата окончания не может быть раньше даты начала."),
        ("Цель", "Кратко и конкретно опишите цель и задачи поездки."),
    ], [2700, 6660])
    add_heading(doc, "Как завершить форму", 2)
    add_number(doc, "Проверьте маршрут, транспорт, даты и цель.")
    add_number(doc, "Если заявка ещё не готова, выберите «Сохранить черновик». Черновик не направляется руководителю.")
    add_number(doc, "Чтобы запустить процесс согласования, выберите «Отправить на согласование».")
    doc.add_paragraph("Приложение показывает предварительный расчёт суточных. Он зависит от количества ночей и установленной в системе ставки. Однодневная поездка не содержит ночей.")
    add_note(doc, "Нерабочие дни", "Субботы, воскресенья и даты из справочника праздников отмечены мягкой розовой заливкой. Их можно выбрать, но приложение попросит подтвердить корректность такой поездки.")

    add_heading(doc, "3. Статусы и согласование", 1)
    doc.add_paragraph("После отправки заявка проходит последовательность, заданную в организационной структуре. Конкретное число этапов зависит от роли сотрудника и цепочки руководителей.")
    add_table(doc, ["Статус", "Значение"], [
        ("Черновик", "Заявка сохранена у сотрудника, но не направлена на согласование."),
        ("На согласовании", "Заявка ожидает действия следующего согласующего."),
        ("Согласовано менеджером", "Менеджер согласовал заявку; при наличии следующего этапа она продолжает движение."),
        ("Согласовано руководителем", "Руководитель согласовал заявку; при необходимости она ожидает финального решения."),
        ("Согласовано", "Заявка полностью утверждена."),
        ("Отклонено", "Один из согласующих отклонил заявку. Проверьте комментарий и при необходимости создайте новую или исправленную заявку."),
    ], [3000, 6360])
    add_heading(doc, "Для руководителя", 2)
    doc.add_paragraph("Откройте «Согласование», выберите заявку подчинённого, проверьте даты, маршрут, цель и сведения о сотруднике. Нажмите «Согласовать» или «Отклонить»; при отклонении добавьте понятный комментарий. Решение фиксируется в истории заявки.")
    add_note(doc, "Порядок согласования", "Система использует назначение руководителя в профиле сотрудника и его роль. Не пытайтесь обходить процесс через чат: изменение руководителя или роли выполняет администратор.")

    add_heading(doc, "4. Календарь", 1)
    doc.add_paragraph("Раздел «Календарь» показывает командировки в месячном, недельном и квартальном виде. Можно переключать период стрелками, вернуться к текущей дате и использовать доступные фильтры сотрудников.")
    for item in [
        "Нажмите на день, чтобы открыть список командировок на эту дату.",
        "Цветная плашка внутри дня отражает статус конкретной командировки.",
        "Мягкая розовая заливка означает нерабочий день: выходной или праздник.",
        "Видимость сотрудников и поездок зависит от вашей роли и организационной структуры.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Чат и обращения по работе приложения", 1)
    doc.add_paragraph("Чат предназначен для рабочей переписки внутри доступного круга сотрудников. В списке контактов обычного сотрудника первым всегда находится «Администратор», далее отображаются коллеги из того же отдела и непосредственный руководитель. У администратора показаны остальные пользователи.")
    add_number(doc, "Откройте «Чат» в боковом меню.")
    add_number(doc, "Выберите получателя слева.")
    add_number(doc, "Введите текст в поле сообщения и нажмите «Отправить» под полем ввода.")
    doc.add_paragraph("На дашборде рядом с кнопкой «Чат» красная цифра показывает количество непрочитанных сообщений.")
    add_note(doc, "Техническая поддержка", "Если проблема относится к работе самого приложения, не используйте общий чат. Откройте «Мой профиль» -> «Связь с администратором», укажите тему и опишите проблему; при необходимости приложите снимок экрана.")

    add_heading(doc, "6. Мой профиль", 1)
    doc.add_paragraph("В профиле можно просмотреть персональные данные, сменить пароль и направить обращение администратору. Организационные данные (отдел, должность, руководитель) влияют на согласование и контакты. Если они неверны, сообщите об этом через форму связи с администратором.")
    add_bullet(doc, "При смене пароля используйте надёжную комбинацию и не сохраняйте её в общем доступе.")
    add_bullet(doc, "После изменения организационных данных администратором обновите страницу, чтобы увидеть актуальные списки и маршрутизацию.")

    add_heading(doc, "7. Если что-то не работает", 1)
    add_table(doc, ["Ситуация", "Что сделать"], [
        ("Не удаётся войти", "Проверьте email и пароль. Если не помогает, направьте обращение через установленный канал администратора."),
        ("Нет нужного маршрута", "Не создавайте маршрут вручную в заявке. Передайте администратору название и последовательность городов для добавления в справочник."),
        ("Нет руководителя или коллег в чате", "Проверьте данные профиля и сообщите администратору ФИО, отдел и руководителя, которые должны быть назначены."),
        ("Заявка не поступила на согласование", "Проверьте, что выбран «Отправить на согласование», а не сохранение черновика. Затем проверьте назначенного руководителя в профиле."),
        ("Не отображаются изменения", "Обновите страницу. При работе на компьютере можно использовать Ctrl+F5 для обновления без старого кеша."),
    ], [3000, 6360])

    add_heading(doc, "8. Ежедневная памятка", 1)
    for item in [
        "Перед отправкой убедитесь, что выбран именно маршрут из списка и корректны даты.",
        "Сохраняйте черновик, если заявка ещё требует уточнений; на согласование отправляйте готовую версию.",
        "Регулярно проверяйте статусы в «Моих командировках».",
        "Руководителям: своевременно просматривайте раздел «Согласование».",
        "Для рабочих вопросов используйте чат, для ошибок и настройки приложения - «Мой профиль» -> «Связь с администратором».",
    ]:
        add_bullet(doc, item)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
