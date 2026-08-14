from pathlib import Path

from docx import Document


for path in sorted(Path("attached_assets").glob("*.docx")):
    print(f"FILE: {path.name}")
    document = Document(path)
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"P{index}: {text}")
    for table_index, table in enumerate(document.tables):
        print(f"TABLE {table_index}")
        for row in table.rows:
            print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
    print("---")
